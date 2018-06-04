# -*- coding: utf-8 -*-
"""
Created on Fri Aug 18 22:52:23 2017

@author: rohit
"""
import pandas
import functools
import datetime
import os
import docx
from docx.shared import Pt
import re

columns_names = pandas.Index(['TRADE_DATE', 'SCRIPCODE', 'PANNO', 'CLIENTNAME', 'CP_PANNO',
            'CP_CLIENTNAME', 'ORDER_NO', 'CP_ORDER_NO', 'TRADEID', 'ORDER_TIME',
            'CP_ORDER_TIME', 'TRADE_TIME', 'ORDER_LMTIME', 'CP_ORDER_LMTIME',
            'TRADE_RATE', 'LTP_RATE', 'LTP_PERCENT', 'TRADE_VALUE', 'ORDER_RATE',
            'CP_ORDER_RATE', 'ORDER_LMRATE', 'CP_ORDER_LMRATE', 'ORDER_LTP',
            'CP_ORDER_LTP', 'TRADED_QTY', 'ORDER_QTY', 'CP_ORDER_QTY',
            'ORDER_AVLQTY', 'CP_ORDER_AVLQTY', 'ORDER_LMQTY', 'CP_ORDER_LMQTY',
            'MEMBER_NO', 'MEMBER_NAME', 'OLD_CLIENT', 'MODIFIED_CLIENT',
            'CP_MEMBER_NO', 'CP_MEMBER_NAME', 'CP_OLD_CLIENT', 'CP_MODIFIED_CLIENT',
            'TERMINALID', 'USER_NAME', 'LOCATION_ADDRESS', 'CP_TEMINALID',
            'CP_USER_NAME', 'CP_LOCATION_ADDRESS', 'ORDER_AUD', 'CP_ORDER_AUD',
            'ORDER_TYPE', 'CP_ORDER_TYPE', 'LOCATIONID', 'CP_LOCATIONID',
            'SCRIPNAME', 'SERIES', 'MKT_TYPE', 'ISIN_NO'],
            dtype='object')

def readfile(path = 'SHREETRD.TXT', sep ='|' ) :
    if '.xls' in path:
        tl = pandas.read_excel(path)
        tl.columns = columns_names
        try:
            tl = formatcsv(tl)
        except:
            raise Exception('bad format')
    else:
        tl = pandas.read_csv(path, sep = '|')
        try:
            tl.columns = columns_names
        except:
            print (tl.columns)
            raise Exception('mismatch in number of columns')
        try:
            tl = formatcsv(tl)
        except:
            raise Exception('bad format')
    tl['HP'] = tl['TRADE_RATE'].cummax()
    tl['NHP'] = tl['HP'] - tl['HP'].shift(1).fillna(tl['HP'].iloc[0])
    tl['LP'] = tl['TRADE_RATE'].cummin()
    tl['NLP'] = tl['LP'] - tl['LP'].shift(1).fillna(tl['LP'].iloc[0])
    return tl

def readfolder(path=r'D:\Python\ricoh', sep ='|') :
    tl = pandas.DataFrame()
    for file in os.listdir(path):
        dest= path + r'\\' + file
        try:
            df = readfile(dest, sep)
            tl = tl.append(df, ignore_index = True)
            print ("loaded ", dest)
        except:
            print("Could not load", dest)
    return tl

def returnposcp(tl,PAN):
    return tl[(tl['PANNO']==PAN)&(tl['LTP_RATE']>0)][['TRADE_DATE','CP_PANNO','CP_CLIENTNAME','LTP_RATE']]

def returncp(tl,PAN):
    return tl[tl['PANNO']==PAN][['TRADE_DATE','CP_PANNO','CP_CLIENTNAME','LTP_RATE']]


def xltp_contributors(tl):
    two= tl.groupby('PANNO').sum()[['LTP_RATE','TRADED_QTY']]
    three= tl.groupby('PANNO').count()['TRADED_QTY']
    four= pandas.merge(two, three.to_frame(), left_index=True, right_index=True).sort_values(by = 'LTP_RATE',ascending= False)
    four.columns = ['LTP_RATE','TRADED_QTY','NO_TRADES']
    return four.sort_values(by = 'NO_TRADES',ascending= False)

def sltp_contributors(tl):
    two= tl.groupby('CP_PANNO').sum()[['LTP_RATE','TRADED_QTY']]
    three= tl.groupby('CP_PANNO').count()['TRADED_QTY']
    four= pandas.merge(two, three.to_frame(), left_index=True, right_index=True).sort_values(by = 'LTP_RATE',ascending= False)
    four.columns = ['LTP_RATE','TRADED_QTY','NO_TRADES']
    return four.sort_values(by = 'NO_TRADES',ascending= True)

def ltp(tl):
    All = xltp_contributors(tl)
    All.rename(columns=lambda x: 'all_'+x, inplace=True)
    pos = xltp_contributors(tl[tl['LTP_RATE']>0])
    pos.rename(columns=lambda x: 'pos_'+x, inplace=True)
    neg = xltp_contributors(tl[tl['LTP_RATE']<0])
    neg.rename(columns=lambda x: 'neg_'+x, inplace=True)
    zero = xltp_contributors(tl[tl['LTP_RATE']==0])
    zero.rename(columns=lambda x: 'zero_'+x, inplace=True)
    zero.drop('zero_LTP_RATE',1, inplace = True)
    dfs = [All, pos, neg, zero]
    total = functools.reduce(lambda left, right: pandas.merge(left, right, left_index=True, right_index=True, how ='outer'), dfs)
    total = total.sort_values(by = 'all_LTP_RATE',ascending= False).fillna(0)
    total['%ltp'] =total['pos_LTP_RATE']*100/total['pos_LTP_RATE'].sum()
    return total

def sltp(tl):
    All = sltp_contributors(tl)
    All.rename(columns=lambda x: 'all_'+x, inplace=True)
    pos = sltp_contributors(tl[tl['LTP_RATE']>0])
    pos.rename(columns=lambda x: 'pos_'+x, inplace=True)
    neg = sltp_contributors(tl[tl['LTP_RATE']<0])
    neg.rename(columns=lambda x: 'neg_'+x, inplace=True)
    zero = sltp_contributors(tl[tl['LTP_RATE']==0])
    zero.rename(columns=lambda x: 'zero_'+x, inplace=True)
    zero.drop('zero_LTP_RATE',1, inplace = True)
    dfs = [All, pos, neg, zero]
    total = functools.reduce(lambda left, right: pandas.merge(left, right, left_index=True, right_index=True, how ='outer'), dfs)
    total = total.sort_values(by = 'all_LTP_RATE',ascending= True).fillna(0)
    total['%ltp'] =total['neg_LTP_RATE']*100/total['neg_LTP_RATE'].sum()
    return total

def top10table(tl,f):
    lp = f(tl)
    top10 = lp[:10]
    table = addtotal(top10)
    table = addmarketotal (lp, table)
    table = addname(tl,table)
    return formatable(table)

def addmarketotal (alltable, table):
    market = alltable.sum()
    market.name = 'Market total'
    table = table.append(market)
    return table

def formatcsv(tl):
    date = tl['TRADE_DATE']
    if type(date.iloc[0]) == str:
        date = date.str.replace('-','/')
        date = date.str.split('/')
        tl['TRADE_DATE'] = date.apply(lambda x: datetime.date(int(x[2]),int(x[1]),int(x[0])))
    a =['ORDER_TIME','CP_ORDER_TIME', 'TRADE_TIME', 'ORDER_LMTIME', 'CP_ORDER_LMTIME']
    for b in a:
        TIME =tl[b]
        if type(TIME.iloc[0]) == str:
            TIME = TIME.str.split(':')
            TIME=TIME.apply(lambda x: datetime.time(int(x[0]), int(x[1]),int(float(x[2])),int((float(x[2])%1)*1000000)))
            tl[b]= TIME
    return tl

def sellordernumber(orderlog, suspect):
    for order in suspect.values:
        date = orderlog[orderlog['ORDERNO']==order[0]]['ORDER_DATE'].values[0]
        number = len(orderlog[(orderlog['ORDER_DATE']==date)&(orderlog['BUYSELL']=='S')])
        print( order[0],'date is ', date, ' : no of sell orders is',number)

def dateformat(date):
    date = date.str.split('/')
    return date.apply(lambda x: datetime.date(int(x[2]),int(x[1]),int(x[0])))

def timeformat(TIME): #h:m:s:ms
    TIME = TIME.str.split(':')
    return TIME.apply(lambda x: datetime.time(int(x[0]), int(x[1]),int(float(x[2])),int((float(x[2])%1)*1000000)))

def hmstimeformat(TIME): #h:m:s
    TIME = TIME.str.split(':')
    return TIME.apply(lambda x: datetime.time(int(x[0]), int(x[1]),int(x[2])))

def suspectable (tl, f):
    suspect = pandas.read_csv('suspect.txt')
    table = f(tl)
    name = table.index.name
    suspecttable = table.reset_index().merge(suspect, left_on = name, right_on= "PAN", how= 'right', ).set_index(name)
    suspecttable = suspecttable.set_index("PAN")
    suspecttable = suspecttable.dropna(how ='all')
    suspecttable = addtotal(suspecttable)
    suspecttable = addmarketotal(table,suspecttable)
    suspecttable = addname (tl, suspecttable)
    return formatable(suspecttable)

def washtrades (tl):
    return tl[tl['PANNO']==tl['CP_PANNO']]

def wash (tl):
    selff = washtrades(tl)
    sel = selff.groupby('PANNO')
    table =xltp_contributors(selff)
    days = sel['TRADE_DATE'].nunique()
    table['NUMBER_DAYS']= days
    posltp = selff[selff['LTP_RATE']>0].groupby('PANNO').LTP_RATE.sum()
    table['POS_LTP'] = posltp
    table['%Market_Volume']=table['TRADED_QTY']/tl['TRADED_QTY'].sum()
    table =table.fillna(0)
    table = table[table.columns[[1,2,3,0,4,5]]]
    return table

def addtotal(table):
    tsum = table.sum()
    tsum.name = 'Total'
    table = table.append(tsum)
    return table

def addname (tl,table):
    b = tl[['PANNO','CLIENTNAME']]
    b = b.drop_duplicates('PANNO')
    b = b.set_index('PANNO')
    c = tl[['CP_PANNO','CP_CLIENTNAME']]
    c = c.drop_duplicates('CP_PANNO')
    c = c.set_index('CP_PANNO')
    c.columns = ['CLIENTNAME']
    b = b.append(c).drop_duplicates()
    name = table.index.name
    table = b.merge(table.reset_index(), left_index = True, right_on= name, how= 'right', ).set_index(name)
    return table

def timediff(a, b):
    import datetime
    # Create datetime objects for each time (a and b)
    dateTimeA = datetime.datetime.combine(datetime.date.today(), a)
    dateTimeB = datetime.datetime.combine(datetime.date.today(), b)
    # Get the difference between datetimes (as timedelta)
    dateTimeDifference = dateTimeA - dateTimeB
    # Divide difference in seconds by number of seconds in hour (3600)  
    dateTimeDifference = abs(dateTimeDifference.total_seconds())
    return dateTimeDifference

def topatch(tl,year,month,day):
    return tl[tl['TRADE_DATE']<=datetime.date(year,month,day)]


def frompatch(tl,year,month,day):
    return tl[tl['TRADE_DATE']>=datetime.date(year,month,day)]

def synctrades (tl):
    ordertime = tl['ORDER_LMTIME'].apply(lambda x: datetime.datetime.combine(datetime.date.today(), x))
    cpordertime = tl['CP_ORDER_LMTIME'].apply(lambda x: datetime.datetime.combine(datetime.date.today(), x))
    timediff = abs (ordertime-cpordertime)
    bt = timediff.apply(lambda x: x.total_seconds())
    sync = tl[(tl['ORDER_LMRATE']==tl['CP_ORDER_LMRATE'])&(tl['ORDER_LMQTY']==tl['CP_ORDER_LMQTY'])&(bt<60)]
    return sync


def sync(tl):
    suspect = pandas.read_csv('suspect.txt')
    suspect = suspect.set_index('PAN')
    suspect['gross buy'] =tl.groupby('PANNO').sum()['TRADED_QTY']
    suspect['gross sell'] =tl.groupby('CP_PANNO').sum()['TRADED_QTY']
    suspect['gross total'] = suspect['gross buy'] + suspect['gross sell']
    among = tradeamongroup(tl)
    suspect['total traded qty among suspected entities'] =among.groupby('PANNO').sum()['TRADED_QTY']
    synchronized = synctrades(among)
    suspect['synchronized traded qty among suspected entities'] =synchronized.groupby('PANNO').sum()['TRADED_QTY']
    suspect['Sync Trades as % of total traded qty among the suspected entities'] = suspect['synchronized traded qty among suspected entities']/suspect['total traded qty among suspected entities'].sum()
    suspect['Sync Trades as % of total market volume'] = suspect['synchronized traded qty among suspected entities']/tl['TRADED_QTY'].sum()
    suspect['Sum of LTP of Sync Trades'] =synchronized.groupby('PANNO').sum()['LTP_RATE']
    suspect['Sum of NHP of Sync Trades'] =synchronized.groupby('PANNO').sum()['NHP']
    suspect = suspect.dropna(how ='all').fillna(0)
    suspect = addtotal(suspect)
    return suspect


def firstrades (tl) :
    trades = tl[tl['TRADE_DATE'] != tl['TRADE_DATE'].shift(1)]
    return trades

def firstbuy(tl):
    ft = firstrades(tl)
    table =pandas.DataFrame([])
    table['Total number of first trades'] = ft.groupby('PANNO').count()['TRADED_QTY']
    table['number of first trades at positive LTP'] = ft[ft['LTP_RATE']>0].groupby('PANNO').count()['TRADED_QTY']
    table['net LTP'] = ft.groupby('PANNO').sum()['LTP_RATE']
    table['Positive LTP'] = ft[ft['LTP_RATE']>0].groupby('PANNO').sum()['LTP_RATE']
    table['NHP'] = ft[ft['NHP']>0].groupby('PANNO').sum()['NHP']
    table = table.fillna(0)
    return table.sort_values(by = 'Total number of first trades',ascending= False)

def firstsell(tl):
    ft = firstrades(tl)
    table =pandas.DataFrame([])
    table['Total number of first trades'] = ft.groupby('CP_PANNO').count()['TRADED_QTY']
    table['number of first trades at negative LTP'] = ft[ft['LTP_RATE']<0].groupby('CP_PANNO').count()['TRADED_QTY']
    table['net LTP'] = ft.groupby('CP_PANNO').sum()['LTP_RATE']
    table['Negative LTP'] = ft[ft['LTP_RATE']<0].groupby('PANNO').sum()['LTP_RATE']
    table = table.fillna(0)
    return table.sort_values(by = 'Total number of first trades',ascending= False)


def nhp (tl):
    nhp = tl[tl['NHP']>0]
    table =pandas.DataFrame([])
    table['Quantity'] =nhp.groupby('PANNO').sum()['TRADED_QTY']
    table['Number of Trades'] =nhp.groupby('PANNO').count()['TRADED_QTY']
    table['Contribution to market NHP'] =nhp.groupby('PANNO').sum()['NHP']
    table['% of market NHP'] = table['Contribution to market NHP']*100/nhp['NHP'].sum()
    return table.sort_values(by = 'Number of Trades',ascending= False)

def nlp (tl):
    nlp = tl[tl['NLP']<0]
    table =pandas.DataFrame([])
    table['Quantity'] =nlp.groupby('CP_PANNO').sum()['TRADED_QTY']
    table['Number of Trades'] =nlp.groupby('CP_PANNO').count()['TRADED_QTY']
    table['Contribution to market NLP'] =nlp.groupby('CP_PANNO').sum()['NLP']
    table['% of market NLP'] = table['Contribution to market NLP']*100/nhp['NLP'].sum()
    return table.sort_values(by = 'Number of Trades',ascending= False)


def formatable(table):
    table= table.fillna('')
    format_mapping={'Float': '{:,.2f}', 'Int': '{:,.0f}', 'Rate': '{:.2f}%'}
    for c in table.columns:
        if 'LTP' in c or 'Contribution to market NHP' in c:
            table[c] = table[c].apply(format_mapping['Float'].format)
        if 'TRADE' in c or 'Quantity' in c or 'Number of Trades' in c:
            table[c] = table[c].apply(format_mapping['Int'].format)
        if '%' in c:
            table[c] = table[c].apply(format_mapping['Rate'].format)
    return table

    
def top10buybrokers (tl):
    brokers= tl.groupby('MEMBER_NO').sum()['TRADED_QTY']
    brokers = brokers.sort_values(ascending= False)
    top10= brokers[:10]
    top10sum = pandas.Series (data = top10.sum(), index = ['Total'])
    market = pandas. Series(brokers.sum(),index=['Market total'])
    tq= brokers.sum()
    table = top10.append(top10sum).append(market)
    table.name = 'TRADED_QTY'
    table =table.to_frame()
    table ['%TQ'] = table['TRADED_QTY']*100/tq
    b = tl[['MEMBER_NO','MEMBER_NAME']]
    b = b.drop_duplicates('MEMBER_NO')
    b = b.set_index('MEMBER_NO')
    table = pandas.merge(b,table, left_index = True, right_index=True, how= 'right')
    return formatable(table)

def top10sellbrokers (tl):
    brokers= tl.groupby('CP_MEMBER_NO').sum()['TRADED_QTY']
    brokers = brokers.sort_values(ascending= False)
    top10= brokers[:10]
    top10sum = pandas.Series (data = top10.sum(), index = ['Total'])
    market = pandas. Series(brokers.sum(),index=['Market total'])
    tq= brokers.sum()
    table = top10.append(top10sum).append(market)
    table.name = 'TRADED_QTY'
    table =table.to_frame()
    table ['%TQ'] = table['TRADED_QTY']*100/tq
    b = tl[['CP_MEMBER_NO','CP_MEMBER_NAME']]
    b = b.drop_duplicates('CP_MEMBER_NO')
    b = b.set_index('CP_MEMBER_NO')
    table = pandas.merge(b,table, left_index = True, right_index=True, how= 'right')
    return formatable(table.fillna(0))


def buyers (tl):
    buyers= tl.groupby('PANNO').sum()['TRADED_QTY']
    table =buyers.to_frame()
    table ['%TQ'] = table['TRADED_QTY']*100/ table['TRADED_QTY'].sum()
    return table.sort_values(by = '%TQ',ascending= False)

def sellers (tl):
    sellers= tl.groupby('CP_PANNO').sum()['TRADED_QTY']
    table =sellers.to_frame()
    table ['%TQ'] = table['TRADED_QTY']*100/ table['TRADED_QTY'].sum()
    return table.sort_values(by = '%TQ',ascending= False)


def trading (tl):
    buyers= tl.groupby('PANNO').sum()['TRADED_QTY']
    sellers= tl.groupby('CP_PANNO').sum()['TRADED_QTY']
    buyers =buyers.to_frame()
    sellers =sellers.to_frame()
    buyers ['%BuyTQ'] = buyers['TRADED_QTY']*100/buyers['TRADED_QTY'].sum()
    sellers ['%SellTQ'] = sellers['TRADED_QTY']*100/buyers['TRADED_QTY'].sum()
    buyers['BUY_TRADED_QTY'] = buyers['TRADED_QTY']
    sellers['SELL_TRADED_QTY'] = sellers['TRADED_QTY']
    del sellers['TRADED_QTY']
    del buyers['TRADED_QTY']
    table =pandas.merge(buyers,sellers, left_index= True, right_index= True, how ='outer')
    table = table.fillna(0)
    table = table.sort_values(by = 'BUY_TRADED_QTY',ascending= False)
    columns =table.columns.tolist()
    columns = [columns[1] , columns[0] , columns[3] , columns[2]]
    table =table[columns]
    table.index.name ='PANNO'
    return table

    
def tradeamongroup(tl):
    suspect = pandas.read_csv('suspect.txt')
    criteria1 =tl['PANNO'].isin(suspect['PAN'].values.tolist())
    criteria2 =tl['CP_PANNO'].isin(suspect['PAN'].values.tolist())
    tl = tl[criteria1 & criteria2]
    return tl
    

def buyordertable(tl , PAN):
    tl = tl[tl['PANNO']==PAN]
    #b = tl['ORDER_LMTIME']>tl['CP_ORDER_LMTIME']
    #tl= tl[b]
    tl['OLMTIME'] = tl['ORDER_LMTIME']
    orss= tl.groupby(['ORDER_NO','OLMTIME'])
    tables =orss[['ORDER_RATE','CP_ORDER_RATE', 'ORDER_LMQTY']].max()
    tables['CP_ORDER_LMQTY'] =orss['CP_ORDER_LMQTY'].sum()
    tables['LTP_RATE'] =orss['LTP_RATE'].sum()
    tables['NO_TRADES'] =orss['LTP_RATE'].count()
    tables['TRADE_DATE'] =orss['TRADE_DATE'].min()
    tables['min CP_ORDER_RATE'] =orss['CP_ORDER_RATE'].min()
    tables['min CP_ORDER_LMTIME'] =orss['CP_ORDER_LMTIME'].min()
    tables['ORDER_LMTIME'] =orss['ORDER_LMTIME'].min()
    return tables.dropna()

def sellordertable(tl , PAN):
    tl = tl[tl['CP_PANNO']==PAN]
    b = tl['ORDER_LMTIME']<tl['CP_ORDER_LMTIME']
    tl= tl[b]
    tl['CP_OLMTIME'] = tl['CP_ORDER_LMTIME']
    orss= tl.groupby(['CP_ORDER_NO','CP_OLMTIME'])
    tables =orss[['ORDER_RATE','CP_ORDER_RATE', 'CP_ORDER_LMQTY']].max()
    tables['ORDER_LMQTY'] =orss['ORDER_LMQTY'].sum()
    tables['LTP_RATE'] =orss['LTP_RATE'].sum()
    tables['NO_TRADES'] =orss['LTP_RATE'].count()
    tables['TRADE_DATE'] =orss['TRADE_DATE'].min()
    return tables.dropna()
 

def add_table(doc, df, lastrows=2):    
    # add a table to the end and create a reference variable
    # extra row is so we can add the header row
    extra_rows = 2 if len(df.columns) == 13 else 1

    t = doc.add_table(df.shape[0]+extra_rows, df.shape[1])
    
    #format
    t.autofit = True
    t.style = 'Table Grid'

    if len(df.columns) == 13:
        xx = ['Ltp Rate', 'Traded Qty', 'No. of Trades']
        df.columns = ['Client Name'] +xx +xx +xx+ xx[1:3] + ['%Ltp']
        t.cell(0,1).merge(t.cell(0,3)).text = 'All trades'
        t.cell(0,4).merge(t.cell(0,6)).text = 'Positive LTP trades'
        t.cell(0,7).merge(t.cell(0,9)).text = 'Negative LTP trades'
        t.cell(0,10).merge(t.cell(0,11)).text = 'All trades'
        for cell in t.rows[0].cells:
            runs = cell.paragraphs[0].runs
            if len(runs)>0:
                run = runs[0]
                run.font.bold = True
                run.font.size = Pt(8)
                run.font.name = 'ArialNarrow'
                cell.paragraphs[0].line_spacing = 0
        


    # add the header rows.
    for j in range(df.shape[-1]):
        t.cell(extra_rows-1,j).text =  ' '.join(df.columns[j].split('_')).title()
        run = t.cell(extra_rows-1,j).paragraphs[0].runs[0]
        run.font.bold = True
        run.font.all_caps = False
        run.font.size = Pt(8)
        run.font.name = 'ArialNarrow'
        t.cell(extra_rows-1,j).paragraphs[0].line_spacing = 0
    

    # add the rest of the data frame
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            t.cell(i+extra_rows,j).text = str(df.values[i,j]).title()
            run = t.cell(i+extra_rows,j).paragraphs[0].runs[0]
            run.font.all_caps = False
            run.font.size = Pt(8)
            run.font.name = 'ArialNarrow'
            t.cell(i+extra_rows,j).paragraphs[0].line_spacing = 0
        t.rows[i].height =0
    for i in range(lastrows):
        t.cell(df.shape[0]+ extra_rows-2 +i,0).text = df.index[-2 + i]
        t.cell(df.shape[0]+ extra_rows-2+i,0).paragraphs[0].runs[0].font.size = Pt(8)
        
    for i in range(lastrows):
        for j in range(df.shape[-1]):
            t.cell(df.shape[0]-i+extra_rows-1,j).paragraphs[0].runs[0].font.bold = True
            t.cell(df.shape[0]-i+extra_rows-1,j).paragraphs[0].runs[0].font.name = 'ArialNarrow'
    return doc

def write_report(file = 'Frontline.xlsx'):
    tl = readfile(file)
    #tl = topatch(tl, 2013,1,30)
    doc = docx.Document()
    doc.add_heading('Investigation Report', 0)
    heading = 'Top 10 Buyers'
    a = top10table(tl ,buyers)
    comment = f"Top 10 buyers contributed {(a['%TQ'].iloc[-2])} to market volume. {a['CLIENTNAME'].iloc[0].title()} was the top buyer with {(a['%TQ'].iloc[0])} contribution to market volume"
    writesection (doc, a, heading, comment)
    heading = 'Top 10 Sellers'
    a = top10table(tl ,sellers)
    comment =f"Top 10 sellers contributed {(a['%TQ'].iloc[-2])} to market volume. {a['CLIENTNAME'].iloc[0].title()} was the top seller with {(a['%TQ'].iloc[0])} contribution to market volume"
    writesection (doc, a, heading, comment)
    heading = 'Top 10 Buy Broker'
    a = top10buybrokers (tl)
    comment = f"Top 10 buy brokers contributed {(a['%TQ'].iloc[-2])} to market volume. {a['MEMBER_NAME'].iloc[0].title()} was the top buy broker with {(a['%TQ'].iloc[0])} contribution to market volume"
    writesection (doc, a, heading, comment)
    heading = 'Top 10 Sell Broker'
    a = top10sellbrokers (tl)
    comment = f"Top 10 sell brokers contributed {(a['%TQ'].iloc[-2])} to market volume. {a['CP_MEMBER_NAME'].iloc[0].title()} was the top sell broker with {(a['%TQ'].iloc[0])} contribution to market volume"
    writesection (doc, a, heading, comment)
    writetop10ltp (doc, tl)
    heading ='Suspect entities Trading'
    a = suspectable(tl ,trading) 
    comment = f"Suspect entities contributed {a['%BuyTQ'].iloc[-2]} to buy volume and {a['%SellTQ'].iloc[-2]} to sell volume"
    writesection (doc, a, heading, comment)
    heading ='Suspect entities LTP contribution'
    a = suspectable(tl ,ltp) 
    comment = f"Suspect entities contributed Rs.{a['all_LTP_RATE'].iloc[-2]} to net LTP and Rs.{a['pos_LTP_RATE'].iloc[-2]} to positive LTP in {a['pos_NO_TRADES'].iloc[-2]} positive LTP trades"
    writesection (doc, a, heading, comment)
    heading = 'Suspect entities NHP contribution'
    a = suspectable(tl ,nhp)
    open = tl['TRADE_RATE'].iloc[0]
    max = tl['TRADE_RATE'].max()
    np = a['Contribution to market NHP'].iloc[-2]
    pernhp = a['% of market NHP'].iloc[-2]
    comment = f"During this period, price of the scrip moved from open price of Rs.{open} to a high price of Rs.{max} and Rs.{max- open} NHP was created. The group contributed Rs.{np} to market NHP. This was {pernhp} of market NHP"
    writesection (doc, a, heading, comment)
    heading = 'Suspect entities self trades'
    a = suspectable(tl ,wash)
    trade  = a['NO_TRADES'].iloc[-2]
    percent = a['%Market_Volume'].iloc[-2]
    comment = f"Suspect entities contributed {percent} to market volume through {trade} self trades"
    writesection (doc, a, heading, comment)
    heading = 'Suspect entities first trades'
    a = suspectable(tl ,firstbuy)
    trade  = a['Total number of first trades'].iloc[-2]
    LTP = a['Positive LTP'].iloc[-2]
    LP = a['NHP'].iloc[-2]
    comment = f"Suspect entities were buyers in {trade} first trades and contributed Rs.{LTP} to market positive LTP and Rs.{LP} to NHP"
    writesection (doc, a, heading, comment)
    heading = 'Suspect entities synchronized trades'
    a = sync(tl)
    a = addname(tl, a)
    trade  = len(synctrades(tradeamongroup(tl)))
    volume = a['Sync Trades as % of total market volume'].iloc[-1]
    LTP = a['Sum of LTP of Sync Trades'].iloc[-1]
    LP = a['Sum of NHP of Sync Trades'].iloc[-1]
    comment = f"Suspect entities through {trade} synchronized trades among themselves, contributed {volume} to market volume and Rs.{LTP} to net LTP "
    writesection (doc, a, heading, comment, 1)
    doc.save('Investigation_report.docx')


def writetop10ltp (doc, tl):
    doc.add_paragraph()
    doc.add_heading('Top 10 LTP Contributors', 1)
    a = top10table(tl ,ltp)
    text = f"During the pre-split period, the price of the scrip opened at \
Rs.{tl['TRADE_RATE'].iloc[0]} on June 01, 2012 and touched period \
high of Rs.{tl['TRADE_RATE'].max()} and closed at \
Rs.{tl['TRADE_RATE'].iloc[0]}"
    doc.add_paragraph(text)
    doc.add_paragraph()
    df = a.columns
    add_table(doc, a)
    doc.add_paragraph()
    a.columns = df
    text = f"Top 10 LTP contributors contributed Rs.{a['all_LTP_RATE'].iloc[-2]} \
to net LTP and and Rs.{a['pos_LTP_RATE'].iloc[-2]} to total market\
positive LTP. {a['CLIENTNAME'].iloc[0].title()} was the top LTP  contributor \
with contribution of Rs.{a['all_LTP_RATE'].iloc[0]} to net LTP and Rs.\
{a['pos_LTP_RATE'].iloc[0]} to total market positive LTP"
    doc.add_paragraph(text)

def writesection (doc, a, heading, comment, lastrows =2):
    doc.add_heading(heading, 1)
    add_table(doc, a, lastrows)
    doc.add_paragraph(comment)

   

    
